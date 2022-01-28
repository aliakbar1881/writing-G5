from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from src.IO import IO_DIR as IO



def doc_gen(words, out_name):
    out_path = str(IO / out_name)+ '.docx'
    document = Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    for key, value in words.items():
        for j in range(2):
            if j == 0:
                even(document, value.word, font_size=12)
                even(document, value.type_, font_size=9)
                document.add_page_break()
            else:
                even(document, value.explain, font_size=9)
                even(document, value.example, font_size=9)
                document.add_page_break()
    print(out_path)
    document.save(out_path)

def even(document, para, font_size=12, font_name='Calibri'):
    paragraph = document.add_paragraph(para)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)

if __name__ == '__main__':
    print('docing.py is script now')
